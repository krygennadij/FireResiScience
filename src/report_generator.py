from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from io import BytesIO

def create_report(data):
    """
    Generates a DOCX report based on the provided data dictionary.
    """
    doc = Document()

    # Style configuration - Times New Roman 12
    # Update 'Normal' style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Update 'Heading 1' style
    h1_style = doc.styles['Heading 1']
    h1_style.font.name = 'Times New Roman'
    h1_style.font.size = Pt(14)
    h1_style.font.bold = True
    h1_style.paragraph_format.space_after = Pt(12)
    h1_style.font.color.rgb = None # Reset color to black if needed, or keep default

    # Update 'Heading 2' style
    h2_style = doc.styles['Heading 2']
    h2_style.font.name = 'Times New Roman'
    h2_style.font.size = Pt(13)
    h2_style.font.bold = True
    h2_style.paragraph_format.space_after = Pt(10)
    h2_style.font.color.rgb = None

    # Helper to add justified paragraph
    def add_p(text, bold=False):
        p = doc.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if bold:
            p.runs[0].bold = True
        return p

    # 1. Header
    h1 = doc.add_heading('ОТЧЕТ', 0)
    h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Title style tweak
    for run in h1.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = None
    
    p_sub = doc.add_paragraph('по расчету фактического предела огнестойкости стальных конструкций')
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph() # Spacer

    # 2. Object Info
    doc.add_heading('1. Наименование и адрес объекта защиты', level=1)
    add_p(f"Объект: {data.get('object_name', 'Не указано')}")
    add_p(f"Адрес: {data.get('object_address', 'Не указано')}")

    # 3. Object Characteristics
    doc.add_heading('2. Краткая характеристика объекта защиты', level=1)
    desc = data.get('object_desc', 'Стальная строительная конструкция.')
    add_p(f"Описание: {desc}")
    
    cp = data.get('calc_params', {})
    add_p(f"Тип элемента: {cp.get('load_type', '-')}")
    add_p(f"Сечение: {cp.get('section_type', '-')} {cp.get('profile_name', '-')}")
    add_p(f"Марка стали: {cp.get('steel_grade', '-')}")
    add_p(f"Действующая нагрузка: {cp.get('n_load', 0)/1000:.2f} кН")

    # 4. Calculation Method
    doc.add_heading('3. Расчетная модель и метод расчета', level=1)
    add_p(
        "Расчет фактического предела огнестойкости выполнен в соответствии с методикой, "
        "изложенной в СП 16.13330 «Стальные конструкции» (в части прочностного расчета) "
        "и Пособии к СП по определению пределов огнестойкости конструкций."
    )
    add_p(
        "Используемый метод: Расчет критической температуры элемента при заданном уровне нагружения "
        "с последующим определением времени прогрева незащищенной конструкции до этой температуры "
        "в условиях стандартного температурного режима пожара (ГОСТ 30247)."
    )

    # 5. Results & Justification
    doc.add_heading('4. Обоснование фактического предела огнестойкости', level=1)
    res = data.get('results', {})
    geom = res.get('geom_props', {})
    
    # 4.0 Geometric Props (New Section)
    doc.add_heading('4.1. Геометрические характеристики сечения', level=2)
    if geom:
        # Convert area to cm2 for readability
        area_cm2 = geom.get('A', 0) * 1e4
        ix_cm4 = geom.get('Ix', 0) * 1e8
        iy_cm4 = geom.get('Iy', 0) * 1e8
        wx_cm3 = geom.get('Wx', 0) * 1e6
        
        add_p(f"Площадь сечения (A): {area_cm2:.2f} см²")
        add_p(f"Момент инерции (Ix): {ix_cm4:.2f} см⁴")
        add_p(f"Момент инерции (Iy): {iy_cm4:.2f} см⁴")
        add_p(f"Момент сопротивления (Wx): {wx_cm3:.2f} см³")
    else:
        add_p("Данные о геометрических характеристиках отсутствуют.")

    # 4.1 Structural calc
    doc.add_heading('4.2. Результаты статического расчета', level=2)
    add_p(f"Коэффициент использования сечения (Gamma_T): {res.get('gamma_t', 0):.4f}")
    add_p(f"Критическая температура стали (t_cr): {res.get('crit_temp', 0):.1f} °C")
    
    # 4.2 Thermal calc
    doc.add_heading('4.3. Результаты теплотехнического расчета', level=2)
    add_p(f"Приведенная толщина металла (delta_np): {res.get('delta_np', 0):.2f} мм")
    add_p(f"Фактическое время достижения критической температуры: {res.get('limit_time', 0):.1f} мин")
    add_p(f"Фактический предел огнестойкости: {res.get('limit_time_str', '-')}")

    # 6. Conclusion
    doc.add_heading('5. Вывод', level=1)
    
    req_time = data.get('required_fire_res', 0)
    act_time = res.get('limit_time', 0)
    
    conclusion = "СООТВЕТСТВУЕТ" if act_time >= req_time else "НЕ СООТВЕТСТВУЕТ"
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    runner = p.add_run(f"Фактический предел огнестойкости конструкции ({res.get('limit_time', 0):.1f} мин) ")
    runner.font.name = 'Times New Roman'
    runner.font.size = Pt(12)
    
    runner_res = p.add_run(f"{conclusion}")
    runner_res.font.name = 'Times New Roman'
    runner_res.font.size = Pt(12)
    runner_res.bold = True
    
    runner_end = p.add_run(f" требуемому пределу огнестойкости ({req_time} мин).")
    runner_end.font.name = 'Times New Roman'
    runner_end.font.size = Pt(12)

    runner_end.font.size = Pt(12)

    # 7. Appendix A: Graph
    if 'graph_image' in res:
        doc.add_page_break()
        doc.add_heading('Приложение А. График нагрева', level=1)
        try:
             # Add picture
             doc.add_picture(BytesIO(res['graph_image']), width=Inches(6.5))
             last_p = doc.paragraphs[-1] 
             last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
             add_p(f"Ошибка при добавлении графика: {e}")

    return doc
